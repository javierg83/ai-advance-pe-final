#!/usr/bin/env python
"""
Este módulo se encarga de la generación de una orden médica en formato Word,
a partir de la información proporcionada por el usuario y el asistente médico.
"""

import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def agregar_separador(document):
    """
    Agrega un separador visual (una línea de guiones) al documento.
    """
    parrafo_sep = document.add_paragraph()
    run_sep = parrafo_sep.add_run("-" * 80)
    run_sep.font.size = Pt(10)
    parrafo_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER


def parse_respuesta_asistente_medico(respuesta):
    """
    Parsea la respuesta del asistente médico y extrae las 3 secciones:
    
      - Análisis de síntomas y factores del paciente
      - Posibles diagnósticos
      - Recomendaciones
    
    La variable 'respuesta' puede ser un diccionario o una cadena formateada.
    """
    if isinstance(respuesta, dict):
        analisis = respuesta.get("analisis", "")
        diagnostico = respuesta.get("diagnostico", "")
        recomendacion = respuesta.get("recomendacion", "")
    elif isinstance(respuesta, str):
        # Expresión regular para extraer el contenido entre los encabezados.
        pattern = (
            r"\*\*Análisis de\s*síntomas y factores del paciente:\*\*(.*?)"
            r"\*\*Posibles diagnósticos:\*\*(.*?)"
            r"\*\*Recomendaciones:\*\*(.*)"
        )
        match = re.search(pattern, respuesta, re.DOTALL)
        if match:
            analisis = match.group(1).strip()
            diagnostico = match.group(2).strip()
            recomendacion = match.group(3).strip()
        else:
            # Si no se encuentra el patrón esperado, se asigna la cadena completa a analisis
            analisis = respuesta.strip()
            diagnostico = ""
            recomendacion = ""
    else:
        analisis = diagnostico = recomendacion = ""
    
    return analisis, diagnostico, recomendacion


def generar_orden_medica(
    openai_client,
    datos_paciente,
    sintomas,
    respuestas_adicionales,
    base_conocimiento,
    respuesta_asistente_medico,
) -> str:
    """
    Genera una orden médica en formato Word a partir de la información proporcionada.

    Parámetros:
    ----------
    openai_client : 
        objeto cliente (para futuras integraciones).
    datos_paciente : 
        diccionario con los datos del paciente.
    sintomas : 
        cadena de texto con los síntomas ingresados por el paciente.
    respuestas_adicionales : 
        diccionario o lista de diccionarios con pares pregunta-respuesta.
    base_conocimiento : 
        información o referencia (no se utiliza en este ejemplo).
    respuesta_asistente_medico : dict o str
        Puede ser un diccionario con claves 'analisis', 'diagnostico' y 'recomendacion',
        o una cadena formateada con los encabezados:
          **Análisis de síntomas y factores del paciente:**
          **Posibles diagnósticos:**
          **Recomendaciones:**
    
    Retorna
    -------
    str
        Nombre del archivo generado.
    """
    
    document = Document()

    # Configurar fuente base (ej. Arial 12 pt)
    estilo = document.styles["Normal"]
    estilo.font.name = "Arial"
    estilo.font.size = Pt(12)

    # Título principal
    titulo = document.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo.add_run("ORDEN MÉDICA")
    run_titulo.bold = True
    run_titulo.font.size = Pt(16)

    document.add_paragraph()  # Línea en blanco

    # Sección: Datos del Paciente
    encabezado_paciente = document.add_paragraph()
    encabezado_paciente.add_run("Datos del Paciente:\n").bold = True
    encabezado_paciente.add_run(f"Nombre: {datos_paciente.get('nombre', '')}\n")
    encabezado_paciente.add_run(f"Edad: {datos_paciente.get('edad', '')}\n")
    encabezado_paciente.add_run(f"Peso: {datos_paciente.get('peso', '')}\n")
    encabezado_paciente.add_run(f"RUT: {datos_paciente.get('rut', '')}\n")

    agregar_separador(document)

    # Sección: Síntomas y Preguntas Adicionales
    parrafo_sintomas = document.add_paragraph()
    parrafo_sintomas.add_run("Síntomas reportados por el paciente:\n").bold = True
    parrafo_sintomas.add_run(sintomas)

    document.add_paragraph()  # Espacio antes de la tabla

    # Crear la tabla para preguntas adicionales
    tabla = document.add_table(rows=1, cols=2)
    tabla.style = "LightShading-Accent1"

    # Encabezados de la tabla
    hdr_cells = tabla.rows[0].cells
    hdr_cells[0].text = "Pregunta"
    hdr_cells[1].text = "Respuesta"

    # Agregar cada par pregunta-respuesta según el tipo de dato
    if isinstance(respuestas_adicionales, dict):
        for pregunta, respuesta in respuestas_adicionales.items():
            row_cells = tabla.add_row().cells
            row_cells[0].text = pregunta
            row_cells[1].text = respuesta
    elif isinstance(respuestas_adicionales, list):
        for item in respuestas_adicionales:
            pregunta = item.get("pregunta", "")
            respuesta = item.get("respuesta", "")
            row_cells = tabla.add_row().cells
            row_cells[0].text = pregunta
            row_cells[1].text = respuesta
    else:
        row_cells = tabla.add_row().cells
        row_cells[0].text = "No se proporcionaron datos válidos."
        row_cells[1].text = ""

    agregar_separador(document)

    # Sección: Datos del Asistente Médico
    parrafo_asistente = document.add_paragraph()
    parrafo_asistente.add_run("Datos generados por el Asistente Médico:\n").bold = True

    # Extraer las tres partes (analisis, diagnostico y recomendacion)
    analisis_text, diagnostico_text, recomendacion_text = parse_respuesta_asistente_medico(respuesta_asistente_medico)

    # 1. Análisis de síntomas y factores del paciente:
    parrafo_analisis = document.add_paragraph()
    parrafo_analisis.add_run("Análisis de síntomas y factores del paciente:\n").bold = True
    parrafo_analisis.add_run(analisis_text)

    document.add_paragraph()  # Espacio

    # 2. Posibles diagnósticos:
    parrafo_diagnostico = document.add_paragraph()
    parrafo_diagnostico.add_run("Posibles diagnósticos:\n").bold = True
    parrafo_diagnostico.add_run(diagnostico_text)

    document.add_paragraph()  # Espacio

    # 3. Recomendaciones:
    parrafo_recomendacion = document.add_paragraph()
    parrafo_recomendacion.add_run("Recomendaciones:\n").bold = True
    parrafo_recomendacion.add_run(recomendacion_text)

    agregar_separador(document)

    # Guardar el documento
    nombre_archivo = "orden_medica.docx"
    document.save(nombre_archivo)
    print(f"Documento guardado como '{nombre_archivo}'")
    return nombre_archivo


def generar_orden_medica_web(
    openai_client,
    datos_paciente,
    sintomas,
    respuestas_adicionales,
    base_conocimiento,
    respuesta_asistente_medico,
):
    """
    Función para entornos web: genera la orden médica usando la función original y devuelve
    la ruta del archivo generado para que se pueda descargar desde la aplicación web.
    """
    archivo = generar_orden_medica(
        openai_client,
        datos_paciente,
        sintomas,
        respuestas_adicionales,
        base_conocimiento,
        respuesta_asistente_medico,
    )
    return archivo


# Ejemplo de llamada para consola (no se ejecuta en modo web)
if __name__ == "__main__":
    # Simulación de datos de entrada para modo escritorio
    openai_client = None  # Se puede pasar None o un objeto de OpenAI si se requiere
    datos_paciente = {
        "nombre": "Juan Pérez",
        "edad": "35",
        "peso": "80kg",
        "rut": "12.345.678-9",
    }
    sintomas = "El paciente presenta dolor abdominal, fiebre y náuseas."

    # Ejemplo de respuestas adicionales (puede ser un diccionario o lista de diccionarios)
    respuestas_adicionales = [
        {
            "pregunta": "1. ¿Desde hace cuánto tiempo ha estado experimentando este dolor y cómo describiría su intensidad?",
            "respuesta": "2 días, intensidad 7/10",
        },
        {
            "pregunta": "2. ¿El dolor es constante o intermitente?",
            "respuesta": "Constante",
        },
        {
            "pregunta": "3. ¿Ha notado algún cambio en sus hábitos intestinales?",
            "respuesta": "No",
        },
        {
            "pregunta": "4. ¿Ha experimentado otros síntomas asociados, como vómitos o fiebre?",
            "respuesta": "Vómitos leves",
        },
        {
            "pregunta": "5. ¿Tiene antecedentes médicos relevantes?",
            "respuesta": "No",
        },
    ]

    base_conocimiento = "Base de datos de conocimiento médico."  # No se utiliza en este ejemplo

    # Ejemplo de respuesta del asistente médico como cadena formateada
    respuesta_asistente_medico = (
        "**Análisis de síntomas y factores del paciente:**\n\n"
        "- **Dolor de cabeza intenso**: Javier ha estado experimentando un dolor de cabeza intenso durante 2 días, con una intensidad de 7 en una escala del 1 al 10. Esto indica un dolor significativo que puede afectar su calidad de vida.\n\n"
        "- **Desencadenante**: El estrés ha sido identificado como un posible desencadenante del dolor de cabeza, lo que sugiere que podría estar relacionado con una cefalea tensional o una migraña.\n\n"
        "- **Síntomas asociados**: La presencia de náuseas puede ser un indicativo de que el dolor de cabeza es más severo y podría estar relacionado con una migraña, ya que este síntoma es común en este tipo de cefaleas.\n\n"
        "- **Antecedentes familiares y personales**: Javier no tiene antecedentes de migrañas o dolores de cabeza recurrentes, lo que puede ser relevante para el diagnóstico.\n\n"
        "**Posibles diagnósticos:**\n\n"
        "1. **Migraña**: Dada la intensidad del dolor, la duración, la presencia de náuseas y el desencadenante de estrés, es probable que Javier esté experimentando una migraña.\n\n"
        "2. **Cefalea tensional**: También es posible que se trate de una cefalea tensional, especialmente si el estrés es un factor significativo.\n\n"
        "**Recomendaciones:**\n\n"
        "1. **Medidas de cuidado**:\n"
        "   - **Descanso**: Asegúrese de descansar en un ambiente oscuro y tranquilo.\n"
        "   - **Hidratación**: Mantenerse bien hidratado puede ayudar a aliviar el dolor.\n"
        "   - **Compresas frías**: Aplicar compresas frías en la frente puede proporcionar alivio.\n\n"
        "2. **Medicamentos**:\n"
        "   - **Analgésicos de venta libre**: Considerar el uso de analgésicos como el ibuprofeno o el paracetamol, siempre que no tenga contraindicaciones para su uso.\n"
        "   - **Medicamentos específicos para migraña**: Si el dolor persiste, podría ser útil consultar a un médico para evaluar la posibilidad de prescribir medicamentos específicos para migrañas, como triptanes.\n\n"
        "3. **Consulta médica**: Si el dolor de cabeza no mejora con las medidas anteriores o si se presentan síntomas adicionales (como cambios en la visión, debilidad o confusión), es importante buscar atención médica inmediata.\n\n"
        "4. **Manejo del estrés**: Considerar técnicas de manejo del estrés, como la meditación, el yoga o la terapia, para ayudar a prevenir futuros episodios.\n\n"
        "Dado que la certeza del diagnóstico de migraña es alta, se recomienda seguir estas pautas y buscar atención médica si los síntomas no mejoran."
    )

    generar_orden_medica(
        openai_client,
        datos_paciente,
        sintomas,
        respuestas_adicionales,
        base_conocimiento,
        respuesta_asistente_medico,
    )
